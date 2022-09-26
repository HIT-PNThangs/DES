import codecs
import docx
import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog
from docx import Document
from PyGui import Ui_MainWindow
from des import Des


class MainWindow:
    def __init__(self):
        self.main_win = QMainWindow()
        self.uic = Ui_MainWindow()
        self.uic.setupUi(self.main_win)
        self.uic.pushButton.clicked.connect(self.openfile)
        self.uic.maHoa.clicked.connect(self.MaHoa)
        self.uic.giaiMa.clicked.connect(self.GiaiMa)
        self.uic.chiaKhoa.clicked.connect(self.ChiaKhoa)
        self.uic.khoiPhucKhoa.clicked.connect(self.khoiphuckhoa)

    def ChiaKhoa(self):
        key = int(self.uic.inputKhoa1.text())
        p = int(self.uic.soPInput_1.text())
        v1 = int(self.uic.v1.text())
        v2 = int(self.uic.v2.text())
        a1 = int(self.uic.a1.text())
        # key = 12379813738877118345
        # v1 = 151595058245452
        # v2 = 111350135012507
        # p = 12764787846358441471
        # a1 = 207244959855905
        fv1 = (a1 * v1 + key) % p
        fv2 = (a1 * v2 + key) % p
        h = 'K1 (v1, f(v1)) : (' + str(v1) + ',' + str(fv1) + ')\nK2 (v2, f(v2)) : (' + str(v2) + ',' + str(fv2) + ')'
        self.uic.textEdit_5.setText(h)
        self.uic.soPInput_2.setText(str(p))
        self.uic.v1_fv1_1.setText(str(v1))
        self.uic.v1_f1_2.setText(str(fv1))
        self.uic.v2_fv2_1.setText(str(v2))
        self.uic.v2_fv2_2.setText(str(fv2))

    def khoiphuckhoa(self):
        p = int(self.uic.soPInput_2.text())
        v1 = int(self.uic.v1_fv1_1.text())
        fv1 = int(self.uic.v1_f1_2.text())
        v2 = int(self.uic.v2_fv2_1.text())
        fv2 = int(self.uic.v2_fv2_2.text())

        b1 = pow(v2 - v1, -1, p) * v2 % p
        b2 = pow(v1 - v2, -1, p) * v1 % p

        key = (fv1 * b1 + fv2 * b2) % p
        self.uic.khoaOutput.setText(str(key))
        key = hex(key)
        self.uic.hex.setText(str(key))

    def show(self):
        self.main_win.show()

    def MaHoa(self):
        new_path = self.uic.editFile.text()
        text = self.uic.textInput.toPlainText()
        key = self.uic.khoaInput.text()
        text = text.encode('utf-8').hex().upper()
        text = text + '0' * (16 - len(text) % 16)
        ListMH = []
        cnt = 0
        while (cnt < len(text) / 16):
            ListMH.append(text[16 * cnt: 16 * (cnt + 1)])
            cnt += 1
        res = ""
        obj = Des()
        for i in range(0, len(ListMH)):
            res += obj.thuchienmahoa(ListMH[i], key)

        self.uic.textOutput.setText(res)
        if len(new_path) != 0:

            if new_path.split('.')[-1] == 'txt':
                with open(new_path, "w") as f:
                    f.write(res)
            else:
                mydoc = Document()
                mydoc.add_paragraph(res)
                mydoc.save(new_path)

    def GiaiMa(self):
        new_path = self.uic.editFile.text()
        text = self.uic.textInput.toPlainText()
        key = self.uic.khoaInput.text()
        ListGM = []
        cnt = 0
        print(len(text))
        print(len(text) / 16)
        while (cnt < len(text) / 16):
            ListGM.append(text[16 * cnt: 16 * (cnt + 1)])
            cnt += 1

        print(len(ListGM))
        res = ""
        obj = Des()
        for i in range(0, len(ListGM)):
            print(ListGM[i])
            print(i)
            res += obj.thuchiengiaima(ListGM[i], key)
        print(res)
        res = codecs.decode(res, 'hex').decode('ASCII')
        res = res[:-1]
        print(res)
        self.uic.textOutput.setText(res)
        if len(new_path) != 0:
            if new_path.split('.')[-1] == 'txt':
                with open(new_path, "w") as f:
                    f.write(res)
            else:
                mydoc = Document()
                mydoc.add_paragraph(res)
                mydoc.save(new_path)

    def openfile(self):
        path = QFileDialog.getOpenFileNames()[0]
        path = path[0]
        self.uic.inputFile.setText(path)
        s = ""
        if path.split('.')[-1] == 'txt':
            with open(path, "r") as f:
                s += f.read()
        else:
            doc = docx.Document(path)
            for i in range(len(doc.paragraphs)):
                s += doc.paragraphs[i].text
                s += '\n'
            s = s[:len(s) - 1]
        self.uic.textInput.setText(s)
        self.uic.khoaInput.setText('133457799BBCDFF1')
        path = path[: path.rfind('.')] + 'new' + path[path.rfind('.'):]
        self.uic.editFile.setText(path)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    main_win = MainWindow()
    main_win.show()
    sys.exit(app.exec())
